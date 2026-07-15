import os
import sys

# ── PDF ───────────────────────────────────────────────────────────────────────
def extract_from_pdf(filepath):
    import pdfplumber

    all_content = []

    with pdfplumber.open(filepath) as pdf:
        total_pages = len(pdf.pages)
        print(f"[PDF] Total pages: {total_pages}")

        for page_num, page in enumerate(pdf.pages, start=1):
            print(f"\n--- Page {page_num} ---")

            # Extract plain text
            text = page.extract_text()
            if text:
                print(text[:500])
                all_content.append({"page": page_num, "type": "text", "content": text.strip()})

            # Extract tables
            tables = page.extract_tables()
            if tables:
                print(f"\n  [Tables found on page {page_num}: {len(tables)}]")
                for t_idx, table in enumerate(tables):
                    print(f"\n  Table {t_idx + 1}:")
                    for row in table:
                        cleaned = [cell if cell else "" for cell in row]
                        print("  | " + " | ".join(cleaned))
                    all_content.append({"page": page_num, "type": "table", "content": table})

    return all_content


# ── EXCEL ─────────────────────────────────────────────────────────────────────
def extract_from_excel(filepath):
    import openpyxl

    wb = openpyxl.load_workbook(filepath)
    all_content = []

    print(f"[Excel] Sheets found: {wb.sheetnames}")

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        print(f"\n--- Sheet: {sheet_name} ---")
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            row_data = [str(cell) if cell is not None else "" for cell in row]
            print("  | " + " | ".join(row_data))
            all_content.append({"sheet": sheet_name, "row": row_idx, "content": row_data})

    return all_content


# ── WORD ──────────────────────────────────────────────────────────────────────
def extract_from_docx(filepath):
    import docx

    doc = docx.Document(filepath)
    all_content = []

    print(f"[DOCX] Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

    # Extract paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"\n[Para {i+1} | {para.style.name}]")
            print(para.text[:300])
            all_content.append({"paragraph": i+1, "style": para.style.name, "type": "text", "content": para.text.strip()})

    # Extract tables
    for t_idx, table in enumerate(doc.tables):
        print(f"\n[Table {t_idx+1}]")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print("  | " + " | ".join(row_data))
        all_content.append({"table_index": t_idx+1, "type": "table", "content": [[cell.text.strip() for cell in row.cells] for row in table.rows]})

    return all_content


# ── TXT / MARKDOWN ────────────────────────────────────────────────────────────
def extract_from_txt(filepath):
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    lines = text.splitlines()
    print(f"[TXT/MD] Total lines: {len(lines)}")
    print("\n--- Preview ---")
    print("\n".join(lines[:50]))

    return [{"type": "text", "content": text}]


# ── MAIN ──────────────────────────────────────────────────────────────────────
def ingest_document(filepath):
    if not os.path.exists(filepath):
        print(f"ERROR: File not found → {filepath}")
        sys.exit(1)

    ext = os.path.splitext(filepath)[1].lower()

    print(f"\n{'='*60}")
    print(f"  File   : {os.path.basename(filepath)}")
    print(f"  Format : {ext}")
    print(f"{'='*60}")

    if ext == ".pdf":
        return extract_from_pdf(filepath)
    elif ext in (".xlsx", ".xls"):
        return extract_from_excel(filepath)
    elif ext in (".docx", ".doc"):
        return extract_from_docx(filepath)
    elif ext in (".txt", ".md"):
        return extract_from_txt(filepath)
    else:
        print(f"Unsupported file type: {ext}")
        sys.exit(1)


if __name__ == "__main__":
    # Change this to test different files
    # Examples:
    #   python 01_ingestion.py sample.pdf
    #   python 01_ingestion.py sample.docx
    #   python 01_ingestion.py sample.xlsx
    #   python 01_ingestion.py sample.txt

    if len(sys.argv) > 1:
        FILE = sys.argv[1]
    else:
        FILE = "sample.pdf"   # default

    extracted = ingest_document(FILE)

    print(f"\n{'='*60}")
    print(f"  Extraction complete.")
    print(f"  Total sections extracted: {len(extracted)}")
    print(f"{'='*60}")